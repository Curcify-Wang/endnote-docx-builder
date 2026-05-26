"""
EndNote-compatible .docx builder template.

WHEN TO USE:
    Whenever the deliverable is a Microsoft Word manuscript whose citations
    will be managed by EndNote. Copy this template into the project folder,
    fill in the REFERENCES list and the manuscript content, and run.

WHY THIS TEMPLATE EXISTS:
    EndNote X9+ does not reliably recognise plain numbered citations [1] or
    temporary citation markers {Author, Year #N}. The only format it accepts
    unconditionally is a real Word complex field whose instruction begins
    with `ADDIN EN.CITE` and contains the EndNote XML. This file generates
    those fields directly.

USAGE:
    1. Copy this file to <project>/build_docx.py
    2. Edit OUT_DIR, FILENAME_PREFIX
    3. Fill in REFERENCES with one entry per cited work
    4. Replace the manuscript content section with the actual paragraphs,
       using {cite(n)} or {cite(n1, n2)} where citations belong
    5. Run: python3 build_docx.py
    6. Inspect post-check output: ADDIN EN.CITE count should equal the
       expected number of citations; suspect {...} fragments must be zero.

NEVER:
    - Write literal {...} text into the manuscript. EndNote CWYW will scan
      ALL visible text and any {Something, Something} substring will trigger
      a "Select Matching Reference" prompt at Update Citations time.
    - Overwrite previous versions. Versioning is auto-incremented.
"""

import os
import re
from datetime import date
from xml.sax.saxutils import escape as xml_escape
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_LINE_SPACING, WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ======================================================================
# CONFIGURATION — edit these
# ======================================================================
OUT_DIR = "/absolute/path/to/your/project/folder"
FILENAME_PREFIX = "Text"       # produces Text<YYYY-MM-DD>_<N>.docx
TODAY = date.today().isoformat()


# ======================================================================
# REFERENCES — one entry per cited work
# ======================================================================
# Each entry needs:
#   temp : "FirstAuthorSurname, Year"  (must match the RIS first AU and PY)
#   ris  : a complete RIS block (multi-line string ending with ER  -)
#
# Two references with identical (Author, Year) are allowed — EndNote
# will prompt to disambiguate on first format. Note these pairs in the
# user-facing README.
#
# The list order determines RecNum (1, 2, 3, ...) used in field codes;
# EndNote will overwrite these on import with the user's library record
# numbers. Library matching is by Author+Year primarily, RecNum
# secondarily.
REFERENCES = [
    # Example entries — replace with real refs for your project
    dict(temp="Baker, 2013",
         ris="""TY  - BOOK
AU  - Baker, R.
TI  - Measuring Walking: A Handbook of Clinical Gait Analysis
CY  - London
PB  - Mac Keith Press
PY  - 2013
ER  -"""),
    dict(temp="Windolf, 2008",
         ris="""TY  - JOUR
AU  - Windolf, M.
AU  - Götzen, N.
AU  - Morlock, M.
TI  - Systematic accuracy and precision analysis of video motion capturing systems—exemplified on the Vicon-460 system
T2  - J Biomech
PY  - 2008
VL  - 41
IS  - 12
SP  - 2776
EP  - 2780
ER  -"""),
    # Add more entries here...
]


# ======================================================================
# Internals — usually no need to edit below this line
# ======================================================================

REF_TYPE_MAP = {
    "JOUR": ("Journal Article", "17"),
    "BOOK": ("Book", "6"),
    "CHAP": ("Book Section", "5"),
    "CONF": ("Conference Proceedings", "10"),
    "COMP": ("Computer Program", "9"),
    "RPRT": ("Report", "27"),
    "THES": ("Thesis", "32"),
    "ELEC": ("Electronic Article", "43"),
}


def next_version_path():
    """Return (path, version_int) for the next Text<DATE>_<N>.docx in OUT_DIR."""
    pat = re.compile(rf"{re.escape(FILENAME_PREFIX)}{TODAY}_(\d+)\.docx$")
    existing = [pat.match(f) for f in os.listdir(OUT_DIR)]
    nums = [int(m.group(1)) for m in existing if m]
    n = (max(nums) + 1) if nums else 1
    return os.path.join(OUT_DIR, f"{FILENAME_PREFIX}{TODAY}_{n}.docx"), n


def parse_ris(ris_text):
    """Parse one RIS block into a dict. AU values accumulate into a list."""
    out = {"AU": []}
    for line in ris_text.strip().split("\n"):
        if "  - " in line:
            tag, value = line.split("  - ", 1)
            tag = tag.strip()
            value = value.strip()
            if tag == "AU":
                out["AU"].append(value)
            else:
                out[tag] = value
    return out


def make_record_xml(ref_index):
    """Build the <record> traveling-library payload for one reference."""
    ref = REFERENCES[ref_index - 1]
    f = parse_ris(ref["ris"])
    ty = f.get("TY", "JOUR")
    type_name, type_num = REF_TYPE_MAP.get(ty, ("Journal Article", "17"))

    x = f"<record><rec-number>{ref_index}</rec-number>"
    x += f'<foreign-keys><key app="EN" db-id="local" timestamp="0">{ref_index}</key></foreign-keys>'
    x += f'<ref-type name="{type_name}">{type_num}</ref-type>'
    x += "<contributors><authors>"
    for au in f["AU"]:
        x += f"<author>{xml_escape(au)}</author>"
    x += "</authors></contributors>"
    if "TI" in f:
        x += f"<titles><title>{xml_escape(f['TI'])}</title>"
        if "T2" in f:
            x += f"<secondary-title>{xml_escape(f['T2'])}</secondary-title>"
        x += "</titles>"
    if "T2" in f:
        x += f"<periodical><full-title>{xml_escape(f['T2'])}</full-title></periodical>"
    pages = ""
    if "SP" in f and "EP" in f:
        pages = f"{f['SP']}-{f['EP']}"
    elif "SP" in f:
        pages = f["SP"]
    if pages:
        x += f"<pages>{xml_escape(pages)}</pages>"
    if "VL" in f:
        x += f"<volume>{xml_escape(f['VL'])}</volume>"
    if "IS" in f:
        x += f"<number>{xml_escape(f['IS'])}</number>"
    if "PY" in f:
        x += f"<dates><year>{xml_escape(f['PY'])}</year></dates>"
    if "PB" in f:
        x += f"<publisher>{xml_escape(f['PB'])}</publisher>"
    if "CY" in f:
        x += f"<pub-location>{xml_escape(f['CY'])}</pub-location>"
    if "DO" in f:
        x += f"<electronic-resource-num>{xml_escape(f['DO'])}</electronic-resource-num>"
    if "UR" in f:
        x += f"<urls><related-urls><url>{xml_escape(f['UR'])}</url></related-urls></urls>"
    x += "</record>"
    return x


def make_endnote_xml(ref_indices):
    """Build the <EndNote><Cite>...</Cite>...</EndNote> XML for one citation field."""
    display = "[" + ", ".join(str(n) for n in ref_indices) + "]"
    x = "<EndNote>"
    for i, n in enumerate(ref_indices):
        ref = REFERENCES[n - 1]
        author, year = (s.strip() for s in ref["temp"].split(",", 1))
        x += "<Cite>"
        x += f"<Author>{xml_escape(author)}</Author>"
        x += f"<Year>{xml_escape(year)}</Year>"
        x += f"<RecNum>{n}</RecNum>"
        if i == 0:
            x += f"<DisplayText>{xml_escape(display)}</DisplayText>"
        x += make_record_xml(n)
        x += "</Cite>"
    x += "</EndNote>"
    return x


def add_endnote_field(paragraph, ref_indices, font_size=12):
    """Inject the five-run Word complex field with ADDIN EN.CITE into a paragraph."""
    display_text = "[" + ", ".join(str(n) for n in ref_indices) + "]"
    field_instr = " ADDIN EN.CITE " + make_endnote_xml(ref_indices) + " "

    r1 = paragraph.add_run()
    fb = OxmlElement("w:fldChar")
    fb.set(qn("w:fldCharType"), "begin")
    r1._r.append(fb)

    r2 = paragraph.add_run()
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = field_instr
    r2._r.append(instr)

    r3 = paragraph.add_run()
    fs = OxmlElement("w:fldChar")
    fs.set(qn("w:fldCharType"), "separate")
    r3._r.append(fs)

    r4 = paragraph.add_run(display_text)
    r4.font.name = "Times New Roman"
    r4.font.size = Pt(font_size)

    r5 = paragraph.add_run()
    fe = OxmlElement("w:fldChar")
    fe.set(qn("w:fldCharType"), "end")
    r5._r.append(fe)


_CITE_PLACEHOLDER = re.compile(r"__CITE_((?:\d+_?)+)__")


def cite(*nums):
    """Return an internal placeholder that populate_paragraph() later turns into
    a real EndNote field. Use inside f-strings:
        f"...prior work {cite(3)} and follow-ups {cite(4, 5, 6)}..."
    """
    return "__CITE_" + "_".join(str(n) for n in nums) + "__"


def populate_paragraph(paragraph, text, font_size=12, italic=False, bold=False):
    """Add text to a paragraph (including table cells), replacing every
    __CITE_n_..._m__ placeholder with a real EndNote Word field. This is
    the ONLY function that should ever add citation-bearing text to the
    document — direct paragraph.add_run(text) will leave placeholders raw."""
    last_end = 0
    for m in _CITE_PLACEHOLDER.finditer(text):
        if m.start() > last_end:
            seg = text[last_end : m.start()]
            r = paragraph.add_run(seg)
            r.font.name = "Times New Roman"
            r.font.size = Pt(font_size)
            r.italic = italic
            r.bold = bold
        nums = [int(n) for n in m.group(1).rstrip("_").split("_") if n]
        add_endnote_field(paragraph, nums, font_size=font_size)
        last_end = m.end()
    if last_end < len(text):
        seg = text[last_end:]
        r = paragraph.add_run(seg)
        r.font.name = "Times New Roman"
        r.font.size = Pt(font_size)
        r.italic = italic
        r.bold = bold


# ======================================================================
# Document scaffolding — typical academic manuscript defaults
# ======================================================================
doc = Document()

for section in doc.sections:
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

style = doc.styles["Normal"]
style.font.name = "Times New Roman"
style.font.size = Pt(12)
rPr = style.element.get_or_add_rPr()
rFonts = rPr.find(qn("w:rFonts"))
if rFonts is None:
    rFonts = OxmlElement("w:rFonts")
    rPr.append(rFonts)
rFonts.set(qn("w:ascii"), "Times New Roman")
rFonts.set(qn("w:hAnsi"), "Times New Roman")
rFonts.set(qn("w:eastAsia"), "Times New Roman")

pf = style.paragraph_format
pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE
pf.first_line_indent = Inches(0.5)


def add_h1(text):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Inches(0)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    r.bold = True
    r.font.name = "Times New Roman"
    r.font.size = Pt(14)
    return p


def add_h2(text):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Inches(0)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.bold = True
    r.font.name = "Times New Roman"
    r.font.size = Pt(12)
    return p


def add_para(text, indent=True, italic_run=False, align=None):
    """Standard manuscript paragraph. Use {cite(n, ...)} inline for citations."""
    p = doc.add_paragraph()
    if not indent:
        p.paragraph_format.first_line_indent = Inches(0)
    if align is not None:
        p.alignment = align
    populate_paragraph(p, text, font_size=12, italic=italic_run)
    return p


# ======================================================================
# MANUSCRIPT CONTENT — replace this section with the actual paper
# ======================================================================
add_h1("Title of the Manuscript")
add_para(
    "Replace this with the actual abstract / introduction. Use "
    f"{cite(1)} for single citations and {cite(1, 2)} for multiple citations. "
    "Never write literal curly braces in user-visible text."
)


# ======================================================================
# Save + RIS bundle + post-check
# ======================================================================
out_path, version_num = next_version_path()
doc.save(out_path)
print(f"Saved manuscript: {out_path}")
print(f"Version: {version_num}")

# RIS bundle for one-shot EndNote library import
ris_path = os.path.join(OUT_DIR, f"references_{TODAY}.ris")
with open(ris_path, "w", encoding="utf-8") as f:
    for r in REFERENCES:
        f.write(r["ris"].strip() + "\n\n")
print(f"Saved RIS bundle: {ris_path}  ({len(REFERENCES)} references)")

# ---- post-check ---------------------------------------------------------
# A. Count real ADDIN EN.CITE Word fields actually emitted
# B. Scan visible text for any leftover {...} patterns — these trigger
#    EndNote's "Select Matching Reference" dialog on Update Citations
import zipfile
with zipfile.ZipFile(out_path) as z:
    document_xml = z.read("word/document.xml").decode("utf-8")
addin_count = document_xml.count("ADDIN EN.CITE")
visible_text = "".join(re.findall(r"<w:t[^>]*>([^<]*)</w:t>", document_xml))
suspect = re.findall(r"\{[^}]*\}", visible_text)

print("---")
print(f"Post-check: ADDIN EN.CITE fields embedded: {addin_count}")
print(f"Post-check: suspect curly-brace text fragments in visible doc: {len(suspect)}")
if suspect:
    print("WARNING — the following visible-text fragments will trigger EndNote prompts:")
    for s in suspect[:10]:
        print("  ->", repr(s))
    print("Remove these literal {...} substrings before delivering the .docx.")
